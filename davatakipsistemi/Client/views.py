from django.contrib import messages
from django.http import HttpResponse
from .models import Client, Case  # Import the Client model
from decimal import Decimal  # Required for Decimal fields
from django.core.paginator import Paginator
from django.contrib.auth.decorators import login_required
from django.template.loader import render_to_string
from weasyprint import HTML
from docx import Document
from django.shortcuts import render, get_object_or_404, redirect


@login_required()
def add_client(request):
    """
    Giriş yapan kullanıcı için yeni bir müvekkil ekler.
    """
    if request.method == 'POST':
        # Formdan gelen veriler
        name = request.POST.get('first_name')
        surname = request.POST.get('last_name')
        tc_no = request.POST.get('tc_no')
        phone = request.POST.get('phone')
        email = request.POST.get('email')
        address = request.POST.get('address')
        
        # Diğer alanlar
        agreement_amount = Decimal(request.POST.get('agreement_amount', '0.00'))
        amount_received = Decimal(request.POST.get('amount_received', '0.00'))
        remaining_balance = agreement_amount - amount_received
        file_expenses = Decimal(request.POST.get('file_expenses', '0.00'))
        
        # Aynı TC numarasıyla giriş yapan kullanıcının bir kaydı var mı kontrol et
        existing_client = Client.objects.filter(tc=tc_no, created_by=request.user).exists()

        if existing_client:
            messages.warning(request, "Bu müvekkil zaten kayıtlı.")
            return redirect('add_client')
        else:
            # Yeni müvekkil kaydı oluştur
            client = Client(
                tc=tc_no,
                name=name,
                surname=surname,
                address=address,
                phone=phone,
                email=email,
                agreement_amount=agreement_amount,
                amount_received=amount_received,
                remaining_balance=remaining_balance,
                file_expenses=file_expenses,
                created_by=request.user  # Giriş yapan kullanıcıyı kaydet
            )
            client.save()
            messages.success(request, "Yeni müvekkil başarıyla eklendi.")
            return redirect(f'/client/{client.id}')

    return render(request, 'client/add_client.html')

@login_required()
def show_client_detail(request, id):
    """
    Giriş yapan kullanıcının belirli bir müvekkilinin detaylarını gösterir.
    """
    # Sadece giriş yapan kullanıcının kayıtlarını getir
    client = get_object_or_404(Client, id=id, created_by=request.user)
    cases_for_client = Case.objects.filter(client_id=id, created_by=request.user)

    context = {
        "client": client,
        "cases_for_client": cases_for_client,
    }
    return render(request, "Client/client.html", context)

@login_required()
def show_client_list(request):
    """
    Giriş yapan kullanıcının müvekkillerini listeler.
    """
    per_page = request.GET.get('per_page', 10)
    sort_by = request.GET.get('sort_by', 'name')
    sort_order = request.GET.get('sort_order', 'asc')
    if sort_order == 'desc':
        sort_by = '-' + sort_by

    # Sadece giriş yapan kullanıcının müvekkilleri
    client_list = Client.objects.filter(created_by=request.user).order_by(sort_by)

    paginator = Paginator(client_list, per_page)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'per_page': per_page,
        'sort_order': sort_order,
        'sort_by': sort_by,
    }
    return render(request, "Client/client_list.html", context)

@login_required()
def edit_client(request, id):
    """
    Giriş yapan kullanıcının belirli bir müvekkilini düzenler.
    """
    client = get_object_or_404(Client, id=id, created_by=request.user)
    
    if request.method == 'POST':
        client.name = request.POST.get('first_name')
        client.surname = request.POST.get('last_name')
        client.tc = request.POST.get('tc_no')
        client.phone = request.POST.get('phone')
        client.email = request.POST.get('email')
        client.address = request.POST.get('address')
        client.agreement_amount = request.POST.get('agreement_amount')
        client.amount_received = request.POST.get('amount_received')
        client.file_expenses = request.POST.get('file_expenses')
        client.save()
        return redirect('client_detail', id=client.id)
    
    return render(request, 'Client/edit_client.html', {'client': client})

@login_required()
def download_client_docx(request, client_id):
    """
    Giriş yapan kullanıcının belirli bir müvekkilinin bilgilerini DOCX formatında indirir.
    """
    client = get_object_or_404(Client, id=client_id, created_by=request.user)
    document = Document()
    
    document.add_heading(f'Müvekkil: {client.name} {client.surname}', 0)
    document.add_paragraph(f'TC Kimlik No: {client.tc}')
    document.add_paragraph(f'Telefon: {client.phone}')
    document.add_paragraph(f'Email: {client.email}')
    document.add_paragraph(f'Anlaşma Miktarı: {client.agreement_amount}')
    # Diğer alanları ekleyin

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{client.name}_{client.surname}.docx"'
    document.save(response)
    return response

@login_required()
def download_client_pdf(request, client_id):
    """
    Giriş yapan kullanıcının belirli bir müvekkilinin bilgilerini PDF formatında indirir.
    """
    client = get_object_or_404(Client, id=client_id, created_by=request.user)
    cases_for_client = Case.objects.filter(client_id=client_id, created_by=request.user)
    
    html_content = render_to_string('Client/client_pdf_template.html', {'client': client, 'cases_for_client': cases_for_client})
    pdf_file = HTML(string=html_content).write_pdf()

    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{client.name}_{client.surname}.pdf"'
    return response
